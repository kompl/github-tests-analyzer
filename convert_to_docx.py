#!/usr/bin/env python3
"""
Конвертация Markdown в красиво отформатированный DOCX
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """Добавить гиперссылку в параграф"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Стиль гиперссылки
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def setup_styles(doc):
    """Настроить стили документа"""
    styles = doc.styles

    # Стиль для кода
    try:
        code_style = styles['Code']
    except KeyError:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)

    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

    # Стиль для inline code
    try:
        inline_code_style = styles['Inline Code']
    except KeyError:
        inline_code_style = styles.add_style('Inline Code', WD_STYLE_TYPE.CHARACTER)

    inline_code_font = inline_code_style.font
    inline_code_font.name = 'Consolas'
    inline_code_font.size = Pt(10)
    inline_code_font.color.rgb = RGBColor(199, 37, 78)

def parse_markdown_line(line):
    """Парсинг элементов форматирования в строке"""
    # Жирный текст
    line = re.sub(r'\*\*(.+?)\*\*', r'{{BOLD}}\1{{/BOLD}}', line)
    # Курсив
    line = re.sub(r'\*(.+?)\*', r'{{ITALIC}}\1{{/ITALIC}}', line)
    # Inline code
    line = re.sub(r'`(.+?)`', r'{{CODE}}\1{{/CODE}}', line)
    # Ссылки
    line = re.sub(r'\[(.+?)\]\((.+?)\)', r'{{LINK}}\1{{SEP}}\2{{/LINK}}', line)

    return line

def apply_formatting(paragraph, text):
    """Применить форматирование к тексту параграфа"""
    parts = re.split(r'({{BOLD}}|{{/BOLD}}|{{ITALIC}}|{{/ITALIC}}|{{CODE}}|{{/CODE}}|{{LINK}}|{{SEP}}|{{/LINK}})', text)

    bold = False
    italic = False
    code = False
    link_text = None
    link_url = None
    in_link = False

    for part in parts:
        if part == '{{BOLD}}':
            bold = True
        elif part == '{{/BOLD}}':
            bold = False
        elif part == '{{ITALIC}}':
            italic = True
        elif part == '{{/ITALIC}}':
            italic = False
        elif part == '{{CODE}}':
            code = True
        elif part == '{{/CODE}}':
            code = False
        elif part == '{{LINK}}':
            in_link = True
            link_text = ''
        elif part == '{{SEP}}':
            link_url = ''
        elif part == '{{/LINK}}':
            if link_text and link_url:
                add_hyperlink(paragraph, link_text, link_url)
            in_link = False
            link_text = None
            link_url = None
        elif part:
            if in_link:
                if link_url is not None:
                    link_url += part
                else:
                    link_text += part
            else:
                run = paragraph.add_run(part)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                if code:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(199, 37, 78)

def convert_markdown_to_docx(md_path, docx_path):
    """Конвертировать Markdown в DOCX с правильным форматированием"""

    doc = Document()
    setup_styles(doc)

    # Установить шрифт по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_block = []
    code_lang = None
    in_table = False
    table_rows = []
    list_level = 0

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Блок кода
        if line.startswith('```'):
            if in_code_block:
                # Конец блока кода
                code_text = '\n'.join(code_block)
                p = doc.add_paragraph(code_text, style='Code')
                p.paragraph_format.left_indent = Inches(0.5)
                # Серый фон
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F5F5F5')
                p._p.get_or_add_pPr().append(shading_elm)

                in_code_block = False
                code_block = []
                code_lang = None
            else:
                # Начало блока кода
                in_code_block = True
                code_lang = line[3:].strip()
            i += 1
            continue

        if in_code_block:
            code_block.append(line)
            i += 1
            continue

        # Заголовки
        if line.startswith('#'):
            match = re.match(r'^(#+)\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                title = match.group(2)

                heading = doc.add_heading(title, level=level)
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Таблицы
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []

            # Убрать разделитель заголовка
            if re.match(r'^\|[\s\-:|]+\|$', line):
                i += 1
                continue

            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_rows.append(cells)

            # Проверить следующую строку
            if i + 1 < len(lines) and '|' not in lines[i + 1]:
                # Конец таблицы
                if table_rows:
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Light Grid Accent 1'

                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            # Убрать markdown форматирование из ячейки
                            clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_data)
                            clean_text = re.sub(r'`(.+?)`', r'\1', clean_text)
                            cell.text = clean_text

                            # Жирный шрифт для первой строки
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True

                in_table = False
                table_rows = []

            i += 1
            continue

        # Списки
        list_match = re.match(r'^(\s*)([-*\d]+\.)\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            marker = list_match.group(2)
            content = list_match.group(3)

            # Определить уровень вложенности
            level = indent // 2

            # Создать параграф списка
            p = doc.add_paragraph(style='List Bullet' if marker in ['-', '*'] else 'List Number')
            p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)

            # Применить форматирование
            formatted_content = parse_markdown_line(content)
            apply_formatting(p, formatted_content)

            i += 1
            continue

        # Обычный текст
        if line.strip():
            p = doc.add_paragraph()
            formatted_line = parse_markdown_line(line)
            apply_formatting(p, formatted_line)
        else:
            # Пустая строка
            doc.add_paragraph()

        i += 1

    # Сохранить документ
    doc.save(docx_path)
    print(f"✅ Документ сохранен: {docx_path}")

if __name__ == '__main__':
    md_path = Path(__file__).parent / 'EXTERNAL_JOBS.md'
    docx_path = Path(__file__).parent / 'EXTERNAL_JOBS_formatted.docx'

    print(f"📄 Конвертирую {md_path} → {docx_path}")
    convert_markdown_to_docx(md_path, docx_path)
